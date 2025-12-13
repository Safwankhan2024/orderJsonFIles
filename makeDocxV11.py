#!/usr/bin/env python3
"""
build_ordered_mcq_book_fixed.py
- Reads order.txt and processes JSON files in that exact order
- Produces one big Markdown book with YAML header + TOC
- Auto-numbers MCQs across chapters
- Also produces a DOCX file with the same content
"""

import json
import os
import glob
import re
from pathlib import Path
from docx import Document
from docx. shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------------
# CONFIG
# --------------------------
INPUT_FOLDER = "mcq_json"            # folder that contains your . json files
ORDER_FILE = "order.txt"             # master order list (one filename per line)
MD_OUTPUT_FILE = "Pediatric_Surgery_MCQ_Book.md"
DOCX_OUTPUT_FILE = "Pediatric_Surgery_MCQ_Book.docx"
BOOK_TITLE = "Pediatric Surgery MCQ Book"
AUTHOR = "Dr.  Safwan Ahmad Khan"
YEAR = "2025"

# show (cleaned) chapter source line under chapter heading?  (True/False)
SHOW_CHAPTER_SOURCE = False

# Answer placement: "immediate", "chapter_end", "book_end"
ANSWER_PLACEMENT = "immediate"

# --------------------------
# YAML metadata header
# --------------------------
YAML_HEADER = f"""---
title: "{BOOK_TITLE}"
author: "{AUTHOR}"
date: "{YEAR}"
toc: true
toc-depth: 2
number-sections: false
description: "A comprehensive collection of pediatric surgery MCQs organized by disease."
---
"""

# --------------------------
# Helpers
# --------------------------

def read_order_file(order_file):
    """Read the order file and return a list of filenames."""
    with open(order_file, 'r') as f:
        lines = [line.strip() for line in f if line.strip()]
    return lines

def find_file_in_input(candidate_name):
    """
    Given candidate from order.txt (with or without .json),
    return the real path in INPUT_FOLDER (case-insensitive match).
    """
    # Simply join with INPUT_FOLDER and check if it exists
    path = os.path.join(INPUT_FOLDER, candidate_name)
    if os.path.exists(path):
        return path
    
    # If not found, try adding .json extension
    if not candidate_name.lower().endswith(".json"):
        path_with_ext = os.path. join(INPUT_FOLDER, candidate_name + ".json")
        if os.path.exists(path_with_ext):
            return path_with_ext
    
    # Try case-insensitive search
    candidate_lower = candidate_name.lower()
    for f in os.listdir(INPUT_FOLDER):
        if f.lower() == candidate_lower:
            return os.path.join(INPUT_FOLDER, f)
    
    return None

def clean_source_display(filename):
    """Remove leading numeric prefix like '74_' and trailing . json/. JSON"""
    name = os.path.basename(filename)
    # strip extension - case-insensitive removal
    if name.lower().endswith(".json"):
        name = name[:-5]
    # remove leading numbers and underscore(s)
    name = re.sub(r"^\d+_+", "", name)
    # Replace underscores with spaces and capitalize
    name = name.replace('_', ' ').replace(' mcqs', '').title()
    return name

def slugify(text):
    """Convert text to a URL-friendly slug."""
    return re.sub(r'[^\w\s-]', '', text).strip().lower().replace(' ', '-')

def load_json_file(file_path):
    """Load JSON file and return data."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='utf-8-sig') as f: # Handle BOM
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"  WARNING:  Skipping malformed JSON in {file_path}")
        return None

def clean_option_text(option_text):
    """Clean option text by removing excessive formatting."""
    # Remove any leading/trailing asterisks or periods
    option_text = re. sub(r'^[\*\.\-\+]+', '', option_text)
    option_text = re.sub(r'[\*\.\-\+]+$', '', option_text)
    # Clean up internal formatting
    option_text = re. sub(r'[\*\._]+', '', option_text)
    return option_text. strip()

# --------------------------
# DOCX Helpers
# --------------------------

def create_docx_document():
    """Create and return a new DOCX document with professional medical book styling."""
    doc = Document()
    
    # Set up styles for professional medical book appearance
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    # Set up heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color. rgb = RGBColor(0x00, 0x00, 0x00)
    
    h3_style = doc. styles['Heading 3']
    h3_style.font. name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    return doc

def add_heading(doc, text, level=1):
    """Add a heading to the DOCX document."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_toc(doc, toc_entries):
    """Add a manual table of contents to the DOCX document."""
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for disease, slug in toc_entries: 
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(f"• {disease}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    doc.add_page_break()

def add_question(doc, mcq_number, question_text):
    """Add a question with professional formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    
    # Question number in bold
    run_num = para.add_run(f"Q {mcq_number}:  ")
    run_num.font. name = 'Times New Roman'
    run_num.font. size = Pt(11)
    run_num.font.bold = True
    
    # Question text in bold
    run_text = para. add_run(question_text)
    run_text.font. name = 'Times New Roman'
    run_text.font. size = Pt(11)
    run_text.font.bold = True

def add_option(doc, option_letter, option_text):
    """Add an option with consistent formatting (no bold)."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_after = Pt(3)
    
    # Option letter and period - NOT bold
    run_letter = para.add_run(f"{option_letter}.  ")
    run_letter.font.name = 'Times New Roman'
    run_letter.font.size = Pt(11)
    run_letter.font. bold = False
    
    # Option text - NOT bold
    run_text = para.add_run(option_text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(11)
    run_text.font. bold = False

def add_answer(doc, answer_text):
    """Add answer with professional formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Inches(0.5)
    
    # "Answer:" label in bold
    run_label = para.add_run("Answer: ")
    run_label.font.name = 'Times New Roman'
    run_label. font.size = Pt(11)
    run_label.font.bold = True
    
    # Answer text in regular font
    run_text = para. add_run(answer_text)
    run_text.font. name = 'Times New Roman'
    run_text.font. size = Pt(11)
    run_text.font.bold = False

def add_explanation(doc, explanation_text):
    """Add explanation with professional formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.left_indent = Inches(0.5)
    
    # "Explanation:" label in bold
    run_label = para. add_run("Explanation: ")
    run_label.font. name = 'Times New Roman'
    run_label.font. size = Pt(11)
    run_label.font.bold = True
    
    # Explanation text in regular font
    run_text = para.add_run(explanation_text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(11)
    run_text.font.bold = False

def add_separator(doc):
    """Add a clean separator line between questions."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run("―" * 60)  # Using a single em-dash character repeated
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# --------------------------
# Build function
# --------------------------

def build_book():
    """Builds the MCQ book from JSON files based on the order specified in ORDER_FILE."""
    os.makedirs(INPUT_FOLDER, exist_ok=True)

    if not os.path.exists(ORDER_FILE):
        raise FileNotFoundError(f"Order file not found: {ORDER_FILE}")

    ordered_filenames = read_order_file(ORDER_FILE)
    print(f"Processing {len(ordered_filenames)} files from {ORDER_FILE}")

    all_chapters = []
    all_chapters_data = []  # Store structured data for DOCX
    toc_entries = []
    mcq_counter = 1
    book_answer_blocks = []
    files_processed = 0
    files_skipped = 0

    for listed in ordered_filenames:
        file_path = find_file_in_input(listed)
        if not file_path:
            print(f"WARNING: Listed file not found: {listed} (skipping)")
            files_skipped += 1
            continue

        data = load_json_file(file_path)
        if data is None:
            files_skipped += 1
            continue

        # Always derive the disease name from the filename for consistent ordering
        disease = clean_source_display(file_path)
        disease = disease.strip()
        chapter_slug = slugify(disease)

        print(f"✓ Processing:  {disease}")
        files_processed += 1

        chapter_content = []
        chapter_answers = []
        chapter_mcqs_data = []  # Store MCQs with their numbers

        chapter_content.append(f"\n\n# {disease}\n")
        if SHOW_CHAPTER_SOURCE: 
            display_src = clean_source_display(os.path.basename(file_path))
            chapter_content.append(f"*Chapter source file: `{display_src}`*\n")

        toc_entries.append((disease, chapter_slug))

        mcqs = data.get("mcqs", [])

        for q in mcqs:
            q_text = q.get("question", "").strip()
            if not q_text:
                print(f"  WARNING:  Skipping malformed MCQ in {file_path}")
                continue

            options = q.get("options", {})
            ans = str(q.get("correct_answer", "")).strip()
            exp = q.get("explanation", "").strip()

            # Store structured data for DOCX
            chapter_mcqs_data.append({
                'number': mcq_counter,
                'question': q_text,
                'options': options,
                'answer': ans,
                'explanation':  exp
            })

            chapter_content.append(f"### MCQ {mcq_counter}:  {q_text}\n")
            # Ensure options are sorted, typically A, B, C, D
            for key in sorted(options.keys()):
                option_text = clean_option_text(options[key])
                chapter_content.append(f"- {key}. {option_text}")

            if ANSWER_PLACEMENT == "immediate": 
                chapter_content.append(f"\n**Answer:** {ans}")
                if exp:
                    chapter_content.append(f"\n**Explanation:** {exp}\n")
                chapter_content.append("\n---\n")
            else:
                chapter_content.append("\n")
                chapter_answers.append((mcq_counter, ans, exp))
                book_answer_blocks.append((mcq_counter, disease, q_text, ans, exp))

            mcq_counter += 1

        if ANSWER_PLACEMENT == "chapter_end" and chapter_answers:
            chapter_content.append("\n**Answer Key (this chapter)**\n")
            for num, ans, exp in chapter_answers:
                chapter_content.append(f"- **MCQ {num}** — **{ans}**:  {exp}")
            chapter_content.append("\n---\n")

        all_chapters.append("".join(chapter_content))
        all_chapters_data.append({
            'disease': disease,
            'mcqs': chapter_mcqs_data
        })

    md_lines = []
    md_lines. append(YAML_HEADER)

    toc_lines = ["## Table of Contents\n"]
    for disease, slug in toc_entries:
        toc_lines.append(f"- [{disease}](#{slug})")

    md_lines.append("\n".join(toc_lines))
    md_lines.append("\n---\n")

    md_lines.extend(all_chapters)

    if ANSWER_PLACEMENT == "book_end" and book_answer_blocks: 
        md_lines.append("\n# Complete Answer Key\n")
        for num, disease, qtext, ans, exp in book_answer_blocks:
            md_lines.append(f"- **MCQ {num}** ({disease}) — **{ans}**: {exp}")

    with open(MD_OUTPUT_FILE, "w", encoding="utf-8") as out:
        out.write("\n".join(md_lines))

    print(f"\n{'='*80}")
    print(f"✓ Markdown output:  {MD_OUTPUT_FILE}")
    print(f"✓ Files processed: {files_processed}")
    print(f"✗ Files skipped: {files_skipped}")
    print(f"✓ Total MCQs: {mcq_counter - 1}")
    print(f"{'='*80}\n")

    # Generate DOCX file with structured data
    build_docx_book(toc_entries, all_chapters_data, book_answer_blocks, mcq_counter)

def build_docx_book(toc_entries, all_chapters_data, book_answer_blocks, mcq_counter):
    """Builds the DOCX book from the processed data with professional medical book formatting."""
    print("Generating DOCX file...")
    
    doc = create_docx_document()

    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Inches(2)
    title_run = title_para.add_run(BOOK_TITLE)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    doc.add_paragraph()  # Space
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(AUTHOR)
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(14)
    
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_run = year_para.add_run(YEAR)
    year_run.font.name = 'Times New Roman'
    year_run.font.size = Pt(12)
    
    doc.add_page_break()

    # Add TOC
    add_toc(doc, toc_entries)

    # Process all chapters with structured data
    chapter_counter = 0
    for chapter_data in all_chapters_data:
        disease = chapter_data['disease']
        mcqs = chapter_data['mcqs']
        
        # Add chapter heading
        if chapter_counter > 0:
            doc.add_page_break()
        
        chapter_heading = doc.add_heading(disease, level=1)
        chapter_counter += 1

        for mcq_data in mcqs:
            mcq_num = mcq_data['number']
            q_text = mcq_data['question']
            options = mcq_data['options']
            ans = mcq_data['answer']
            exp = mcq_data['explanation']
            
            # Add question
            add_question(doc, mcq_num, q_text)
            
            # Add options
            for key in sorted(options.keys()):
                option_text = clean_option_text(options[key])
                add_option(doc, key, option_text)
            
            # Add answer and explanation
            if ANSWER_PLACEMENT == "immediate": 
                add_answer(doc, ans)
                if exp: 
                    add_explanation(doc, exp)
                add_separator(doc)

    doc.save(DOCX_OUTPUT_FILE)
    print(f"✓ DOCX output:  {DOCX_OUTPUT_FILE}")

if __name__ == "__main__":
    build_book()